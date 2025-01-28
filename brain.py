import os
import random
import io
import streamlit as st
import json
import pickle
import numpy as np
from datetime import datetime
from docx import Document as DocxDocument
import re
from keras.models import load_model

# Cargar archivos locales directamente sin usar requests
comands_path = './comands.json'
words_path = './words.pkl'
classes_path = './classes.pkl'
model_path = './chatbot_model.h5'

# Cargar los archivos locales
with open(comands_path) as f:
    intents = json.load(f)

words = pickle.load(open(words_path, 'rb'))
classes = pickle.load(open(classes_path, 'rb'))
model = load_model(model_path)  # Cargar el modelo de la ruta

# Función que tokeniza usando expresiones regulares
def clean_up_sentence(sentence):
    sentence_words = re.findall(r'\b\w+\b', sentence.lower())
    return sentence_words

# Convertimos la información a unos y ceros según si están presentes en los patrones
def bag_of_words(sentence):
    sentence_words = clean_up_sentence(sentence)
    bag = [0] * len(words)
    for w in sentence_words:
        for i, word in enumerate(words):
            if word == w:
                bag[i] = 1
    return np.array(bag)

# Predecimos la categoría a la que pertenece la oración
def predict_class(sentence):
    bow = bag_of_words(sentence)
    res = model.predict(np.array([bow]))[0]
    max_index = np.where(res == np.max(res))[0][0]
    category = classes[max_index]
    return category

# Diccionario de rutas de documentos
doc_paths = {
    "Sistemas automotrices semestre 7-tecnologia de materiales automotrices": './path_to_your_file.docx',
    "Sistemas automotrices semestre 7-programasintetico-tecnologia de materiales automotrices": './path_to_your_file.pdf',
    "Ingenieria mecatronica semestre 1-programasintetico-estructura y propiedades de los materiales": './path_to_your_file.docx',
    "Ingenieria mecatronica semestre 1-estructura y propiedades de los materiales": './path_to_your_file.pdf'
}

# Función para manejar el documento y proporcionar el botón de descarga
def handle_document(tag):
    doc_path = doc_paths.get(tag)

    if not doc_path or not os.path.exists(doc_path):
        return "Documento no encontrado para el semestre o carrera solicitado."

    try:
        if doc_path.endswith('.docx'):
            doc = DocxDocument(doc_path)
            if len(doc.tables) > 0:
                tabla_fecha = doc.tables[0]
                if len(tabla_fecha.rows) > 0 and len(tabla_fecha.columns) > 1:
                    fila_index = 0
                    columna_index = 1
                    celda = tabla_fecha.rows[fila_index].cells[columna_index]
                    current_date = datetime.today().strftime('%m/%d/%y %H:%M:%S')
                    celda.text = f"Fecha: {current_date}"

            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            st.download_button(
                label="Descargar Temario",
                data=buffer,
                file_name=f"Temario_{tag.replace(' ', '_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        elif doc_path.endswith('.pdf'):
            with open(doc_path, "rb") as pdf_file:
                pdf_buffer = pdf_file.read()

            st.download_button(
                label="Descargar programa sintetico",
                data=pdf_buffer,
                file_name=f"Programa_Sintetico_{tag.replace(' ', '_')}.pdf",
                mime="application/pdf"
            )

    except Exception as e:
        st.error(f"Ocurrió un error al manejar el archivo: {e}")
    return ""

# Obtener una respuesta aleatoria
def get_response(tag, intents_json):
    list_of_intents = intents_json['intents']
    result = ""
    for i in list_of_intents:
        if i["tag"] == tag:
            result = random.choice(i['responses'])
            break

    if tag in doc_paths:
        doc_result = handle_document(tag)
        result += "\n" + doc_result  # Agregar resultado del manejo de documento

    return result if result else "No se encontró una respuesta adecuada."
